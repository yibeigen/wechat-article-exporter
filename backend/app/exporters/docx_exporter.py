import io
import re
import html
import datetime
from pathlib import Path
from typing import List, Optional, Dict, Tuple
import httpx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from app.exporters.base import BaseExporter
from app.models import ArticleItem
from app.config import BRAND_OFFICIAL_ACCOUNT, BRAND_FOOTER_NOTE, BRAND_DISCLAIMER

def get_platform_referer(url: str) -> Dict[str, str]:
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
    }
    u = url.lower()
    if "cnblogs" in u: headers["Referer"] = "https://www.cnblogs.com/"
    elif "csdn" in u: headers["Referer"] = "https://blog.csdn.net/"
    elif "zhihu" in u or "zhimg" in u: headers["Referer"] = "https://www.zhihu.com/"
    elif "juejin" in u: headers["Referer"] = "https://juejin.cn/"
    elif "weixin" in u or "qpic" in u: headers["Referer"] = "https://mp.weixin.qq.com/"
    elif "51cto" in u: headers["Referer"] = "https://blog.51cto.com/"
    return headers

def set_cell_background(cell, fill_hex: str):
    """设置单元格背景颜色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    """设置单元格内边距 (1 pt = 20 dxa)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1"):
    """设置标准浅灰细表格边框"""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_callout_borders(table, border_color="3B82F6"):
    """设置 Callout 引用框左侧品牌强调色边框，其余边框隐藏"""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="none"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_codebox_borders(table, border_color="E2E8F0"):
    """设置代码块四周浅灰细边框"""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'<w:insideH w:val="none"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def format_run(
    r,
    font_name="Microsoft YaHei",
    font_size=Pt(10.5),
    color=RGBColor(30, 41, 59),
    bold=False,
    italic=False,
    strike=False,
    underline=False
):
    """为 Run 设置字体（必须同时绑定西文字体与东亚中文字体，彻底解决中文字体粗细不均、大小不一）"""
    r.font.name = font_name
    r._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font_name)
    r.font.size = font_size
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    if strike:
        r.font.strike = True
    if underline:
        r.underline = True

def sanitize_markdown_text(text: str) -> str:
    """预处理并清洗 Markdown/HTML 杂质与转义实体"""
    if not text:
        return ""
    # 1. HTML 实体转义
    t = html.unescape(text)
    # 2. 替换常见残留 HTML 标签
    t = re.sub(r'<br\s*/?>', '\n', t, flags=re.IGNORECASE)
    t = re.sub(r'</?(?:b|strong)>', '**', t, flags=re.IGNORECASE)
    t = re.sub(r'</?(?:i|em)>', '*', t, flags=re.IGNORECASE)
    t = re.sub(r'</?(?:del|s|strike)>', '~~', t, flags=re.IGNORECASE)
    t = re.sub(r'</?code>', '`', t, flags=re.IGNORECASE)
    t = re.sub(r'</?u>', '', t, flags=re.IGNORECASE)
    t = re.sub(r'</?mark>', '', t, flags=re.IGNORECASE)
    # 剥离多余的容器标签
    t = re.sub(r'</?(?:div|p|span|font|section|article)[^>]*>', '', t, flags=re.IGNORECASE)
    # 3. 剥离 Markdown 转义斜杠
    t = re.sub(r'\\([\\`\*_{}\[\]\(\)#\+\-\.!~|])', r'\1', t)
    # 4. 过滤文章内部的目录标记
    if t.strip() in ["@目录", "[TOC]", "目录", "[toc]"]:
        return ""
    return t

def add_inline_markdown_runs(
    paragraph,
    text: str,
    default_font_size=Pt(10.5),
    default_color=RGBColor(30, 41, 59),
    is_bold=False,
    is_italic=False
):
    """
    高精度分词解析行内样式：**粗体**, *斜体*, ***粗斜体***, `代码`, ~~删除线~~, [文本](链接)
    所有分词严格继承父级字号尺寸与字体体系，彻底杜绝字号突变跳跃。
    """
    clean_text = sanitize_markdown_text(text)
    if not clean_text:
        return

    pattern = re.compile(
        r'(`(?P<code>[^`]+)`)'
        r'|(\*\*\*(?P<bold_italic>[^\*]+)\*\*\*)'
        r'|(___(?P<bold_italic2>[^_]+)___)'
        r'|(\*\*(?P<bold>[^\*]+)\*\*)'
        r'|(__(?P<bold2>[^_]+)__)'
        r'|(\*(?P<italic>[^\*]+)\*)'
        r'|(_(?P<italic2>[^_]+)_)'
        r'|(~~(?P<strike>[^~]+)~~)'
        r'|(\[(?P<link_text>[^\]]+)\]\((?P<link_url>[^\)]+)\))'
    )

    last_idx = 0
    for m in pattern.finditer(clean_text):
        start, end = m.span()
        if start > last_idx:
            plain = clean_text[last_idx:start]
            r = paragraph.add_run(plain)
            format_run(r, font_size=default_font_size, color=default_color, bold=is_bold, italic=is_italic)

        if m.group("code"):
            code_str = m.group("code")
            r = paragraph.add_run(code_str)
            # 行内代码：字号微调但不突兀，使用等宽字体 Consolas 与玫红强调色
            code_sz = max(default_font_size - Pt(0.5), Pt(9))
            format_run(r, font_name="Consolas", font_size=code_sz, color=RGBColor(190, 18, 60), bold=is_bold)
        elif m.group("bold_italic") or m.group("bold_italic2"):
            bi_str = m.group("bold_italic") or m.group("bold_italic2")
            r = paragraph.add_run(bi_str)
            format_run(r, font_size=default_font_size, color=RGBColor(15, 23, 42), bold=True, italic=True)
        elif m.group("bold") or m.group("bold2"):
            b_str = m.group("bold") or m.group("bold2")
            r = paragraph.add_run(b_str)
            format_run(r, font_size=default_font_size, color=RGBColor(15, 23, 42), bold=True, italic=is_italic)
        elif m.group("italic") or m.group("italic2"):
            i_str = m.group("italic") or m.group("italic2")
            r = paragraph.add_run(i_str)
            format_run(r, font_size=default_font_size, color=default_color, bold=is_bold, italic=True)
        elif m.group("strike"):
            s_str = m.group("strike")
            r = paragraph.add_run(s_str)
            format_run(r, font_size=default_font_size, color=RGBColor(148, 163, 184), bold=is_bold, italic=is_italic, strike=True)
        elif m.group("link_text"):
            lt = m.group("link_text")
            lu = m.group("link_url")
            # 过滤内部锚点链接（例如 #1-标题），直接展示干净文本
            if lu.startswith("#"):
                r = paragraph.add_run(lt)
                format_run(r, font_size=default_font_size, color=default_color, bold=is_bold, italic=is_italic)
            else:
                if lt == lu:
                    r = paragraph.add_run(lt)
                else:
                    r = paragraph.add_run(f"{lt} ({lu})")
                format_run(r, font_size=default_font_size, color=RGBColor(2, 132, 199), bold=is_bold, italic=is_italic)

        last_idx = end

    if last_idx < len(clean_text):
        plain = clean_text[last_idx:]
        r = paragraph.add_run(plain)
        format_run(r, font_size=default_font_size, color=default_color, bold=is_bold, italic=is_italic)

def render_callout_box(doc: Document, quote_lines: List[str]):
    """渲染专业优雅的 Callout 引用框 (左侧品牌蓝粗边线 + 淡蓝灰底纹 + 舒适内边距)"""
    if not quote_lines:
        return
    
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_callout_borders(table, border_color="3B82F6")
    
    cell = table.cell(0, 0)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=180)
    
    for idx, q_line in enumerate(quote_lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3) if idx < len(quote_lines) - 1 else Pt(0)
        p.paragraph_format.line_spacing = 1.3
        
        cleaned = q_line.strip()
        add_inline_markdown_runs(
            p,
            cleaned,
            default_font_size=Pt(9.5),
            default_color=RGBColor(71, 85, 105),
            is_italic=False
        )

    # 引用框后的缓冲间距
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)

def render_code_box(doc: Document, code_text: str):
    """渲染独立的代码容器块 (等宽字体 Consolas + 细浅灰边框 + 浅灰底纹)"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_codebox_borders(table, border_color="E2E8F0")
    
    cell = table.cell(0, 0)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.18
    
    r = p.add_run(code_text)
    format_run(r, font_name="Consolas", font_size=Pt(9.0), color=RGBColor(30, 41, 59))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(6)

def render_markdown_table(doc: Document, table_lines: List[str]):
    """解析并渲染 Markdown 表格为专业 Word 表格 (表头灰底加粗 + 隔行斑马线交替底色)"""
    rows_data = []
    for line in table_lines:
        line_str = line.strip().strip('|')
        if not line_str or re.match(r'^[\s\-:|]+$', line_str):
            continue
        cells = [c.strip() for c in line_str.split('|')]
        rows_data.append(cells)

    if not rows_data:
        return

    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "CBD5E1")

    for row_idx, row in enumerate(rows_data):
        for col_idx, cell_text in enumerate(row):
            if col_idx >= num_cols:
                continue
            cell = table.cell(row_idx, col_idx)
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15

            if row_idx == 0:
                # 表头：淡蓝灰背景 + 统一字号加粗
                set_cell_background(cell, "F1F5F9")
                add_inline_markdown_runs(p, cell_text, default_font_size=Pt(9.5), default_color=RGBColor(15, 23, 42), is_bold=True)
            else:
                # 表体隔行交替底纹
                if row_idx % 2 == 0:
                    set_cell_background(cell, "F8FAFC")
                else:
                    set_cell_background(cell, "FFFFFF")
                add_inline_markdown_runs(p, cell_text, default_font_size=Pt(9.5), default_color=RGBColor(51, 65, 85))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(6)

def append_article_content_to_docx(
    doc: Document,
    art: ArticleItem,
    author_name: str,
    platform: str,
    embed_images: bool = True
):
    """高质量逐篇渲染排版 (出版级字阶、代码框、Callout 引用框、中文字体严密对齐)"""
    # 1. 文章大标题 (小二 18pt，加粗深蓝)
    h_p = doc.add_paragraph()
    h_p.paragraph_format.space_before = Pt(16)
    h_p.paragraph_format.space_after = Pt(6)
    h_r = h_p.add_run(art.title)
    format_run(h_r, font_size=Pt(18), color=RGBColor(15, 23, 42), bold=True)

    # 2. 元数据说明栏
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(0)
    meta_p.paragraph_format.space_after = Pt(8)
    meta_r = meta_p.add_run(
        f"作者: {art.author or author_name}   |   发布时间: {art.publish_time or '未知'}   |   来源平台: {art.platform or platform}\n"
    )
    format_run(meta_r, font_size=Pt(9), color=RGBColor(100, 116, 139))
    if art.url:
        url_r = meta_p.add_run(f"原文链接: {art.url}")
        format_run(url_r, font_size=Pt(9), color=RGBColor(2, 132, 199))

    # 装饰分割线
    div_p = doc.add_paragraph()
    div_p.paragraph_format.space_after = Pt(12)
    div_r = div_p.add_run("―" * 45)
    format_run(div_r, font_size=Pt(8), color=RGBColor(226, 232, 240))

    # 3. 正文逐行解析引擎
    raw_text = art.content_markdown or ""
    # 统一换行符并分割
    lines = raw_text.replace("\r\n", "\n").split("\n")
    
    in_code_block = False
    code_buffer = []
    quote_buffer = []
    table_buffer = []

    # 图片匹配正则
    img_pattern = re.compile(r"!\[(?P<alt>.*?)\]\((?P<src>https?://[^\s\)]+)\)")

    def flush_quote():
        nonlocal quote_buffer
        if quote_buffer:
            render_callout_box(doc, quote_buffer)
            quote_buffer = []

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            render_markdown_table(doc, table_buffer)
            table_buffer = []

    # 预处理：合并孤立的项目符号行（跳过空行找到对应内容）
    processed_lines = []
    i = 0
    while i < len(lines):
        curr = lines[i].strip()
        if curr in ["•", "●", "·", "▪", "▫", "-", "*", "+"]:
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines):
                next_l = lines[j].strip()
                if next_l and not next_l.startswith(("#", "```", ">", "|")):
                    processed_lines.append(f"• {next_l}")
                    i = j + 1
                    continue
        processed_lines.append(lines[i])
        i += 1

    for line in processed_lines:
        stripped = line.strip()

        # ==========================================================
        # 1. 代码块处理 (```)
        # ==========================================================
        if stripped.startswith("```"):
            flush_quote()
            flush_table()
            if in_code_block:
                code_text = "\n".join(code_buffer)
                if code_text.strip():
                    render_code_box(doc, code_text)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        # ==========================================================
        # 2. 引用块处理 (> ...)
        # ==========================================================
        if stripped.startswith(">"):
            flush_table()
            # 剥离前导 > 并加入 buffer
            content_in_quote = re.sub(r'^>\s?', '', stripped)
            if content_in_quote:
                quote_buffer.append(content_in_quote)
            continue
        else:
            flush_quote()

        # ==========================================================
        # 3. 表格处理 (| ... |)
        # ==========================================================
        if stripped.startswith("|") and ("|" in stripped[1:] or stripped.endswith("|")):
            table_buffer.append(stripped)
            continue
        else:
            flush_table()

        # 忽略纯空行
        if not stripped:
            continue

        # ==========================================================
        # 4. 图片独立行识别与高质量内嵌
        # ==========================================================
        img_match = img_pattern.search(stripped)
        if img_match and (stripped.startswith("![") or stripped.startswith("<img")):
            alt = img_match.group("alt").strip()
            src = img_match.group("src").strip()

            inserted = False
            if embed_images:
                try:
                    headers = get_platform_referer(src)
                    resp = httpx.get(src, headers=headers, timeout=6.0)
                    if resp.status_code == 200 and len(resp.content) > 100:
                        img_bio = io.BytesIO(resp.content)
                        img_p = doc.add_paragraph()
                        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_p.paragraph_format.space_before = Pt(8)
                        img_p.paragraph_format.space_after = Pt(2)
                        img_run = img_p.add_run()
                        img_run.add_picture(img_bio, width=Inches(5.5))
                        inserted = True
                except Exception:
                    pass

            if not inserted:
                fallback_p = doc.add_paragraph()
                fallback_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fallback_p.paragraph_format.space_before = Pt(4)
                fallback_p.paragraph_format.space_after = Pt(4)
                f_run = fallback_p.add_run(f"🖼️ [配图] {alt if alt and alt != '在这里插入图片描述' else ''} ({src})")
                format_run(f_run, font_size=Pt(8.5), color=RGBColor(100, 116, 139), italic=True)
            elif alt and alt != "在这里插入图片描述" and len(alt) < 50:
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.paragraph_format.space_after = Pt(6)
                cap_r = cap_p.add_run(f"▲ {alt}")
                format_run(cap_r, font_size=Pt(8.5), color=RGBColor(148, 163, 184), italic=True)
            continue

        # ==========================================================
        # 5. 分割线 (--- / *** / ___)
        # ==========================================================
        if re.match(r'^(?:-{3,}|\*{3,}|_{3,})$', stripped):
            sep_p = doc.add_paragraph()
            sep_p.paragraph_format.space_before = Pt(6)
            sep_p.paragraph_format.space_after = Pt(8)
            sep_r = sep_p.add_run("――――――――――――――――――――――――――――――――――――――――――")
            format_run(sep_r, font_size=Pt(8), color=RGBColor(226, 232, 240))
            continue

        # ==========================================================
        # 6. 各级标题匹配 (严格字号阶梯，全行格式锁死)
        # ==========================================================
        h_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if h_match:
            hashes, h_text = h_match.group(1), h_match.group(2)
            # 剥离尾部可能存在的 # 标记
            h_text = re.sub(r'\s+#+$', '', h_text).strip()
            level = len(hashes)
            
            h = doc.add_paragraph()
            if level == 1:
                h.paragraph_format.space_before = Pt(14)
                h.paragraph_format.space_after = Pt(4)
                add_inline_markdown_runs(h, h_text, default_font_size=Pt(16), default_color=RGBColor(15, 23, 42), is_bold=True)
            elif level == 2:
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(4)
                add_inline_markdown_runs(h, h_text, default_font_size=Pt(14), default_color=RGBColor(30, 41, 59), is_bold=True)
            elif level == 3:
                h.paragraph_format.space_before = Pt(9)
                h.paragraph_format.space_after = Pt(3)
                add_inline_markdown_runs(h, h_text, default_font_size=Pt(12), default_color=RGBColor(51, 65, 85), is_bold=True)
            elif level == 4:
                h.paragraph_format.space_before = Pt(6)
                h.paragraph_format.space_after = Pt(2)
                add_inline_markdown_runs(h, h_text, default_font_size=Pt(11), default_color=RGBColor(71, 85, 105), is_bold=True)
            else:
                h.paragraph_format.space_before = Pt(5)
                h.paragraph_format.space_after = Pt(2)
                add_inline_markdown_runs(h, h_text, default_font_size=Pt(10.5), default_color=RGBColor(100, 116, 139), is_bold=True)
            continue

        # ==========================================================
        # 7. 列表项匹配 (有序列表 1. / 无序列表 - * + • ● ·)
        # ==========================================================
        # 无序列表 (- , * , + , • , ● , · )
        if re.match(r'^(?:[•●·▪▫\-\*\+])\s+(.*)', stripped):
            item_match = re.match(r'^(?:[•●·▪▫\-\*\+])\s+(.*)', stripped)
            item_text = item_match.group(1) if item_match else stripped
            
            # 计算缩进层级
            leading_spaces = len(line) - len(line.lstrip())
            indent_level = min(leading_spaces // 2, 4)
            
            lp = doc.add_paragraph()
            lp.paragraph_format.left_indent = Inches(0.22 * (indent_level + 1))
            lp.paragraph_format.space_before = Pt(0)
            lp.paragraph_format.space_after = Pt(2.5)
            lp.paragraph_format.line_spacing = 1.3
            
            # 统一圆点符号字体与字号
            bullet_r = lp.add_run("• ")
            format_run(bullet_r, font_size=Pt(10.5), color=RGBColor(37, 99, 235), bold=True)
            
            add_inline_markdown_runs(lp, item_text, default_font_size=Pt(10.5), default_color=RGBColor(30, 41, 59))
            continue

        # 有序列表 (1. , 2. , 12. , (1) , ① )
        if re.match(r'^(?:\(?\d+[\.\、\)]|\([一二三四五六七八九十]+\)|[①②③④⑤⑥⑦⑧⑨⑩])\s*(.*)', stripped):
            num_match = re.match(r'^(\(?\d+[\.\、\)]|\([一二三四五六七八九十]+\)|[①②③④⑤⑥⑦⑧⑨⑩])\s*(.*)', stripped)
            num_prefix = num_match.group(1) if num_match else "1. "
            item_text = num_match.group(2) if num_match else stripped
            
            leading_spaces = len(line) - len(line.lstrip())
            indent_level = min(leading_spaces // 2, 4)
            
            lp = doc.add_paragraph()
            lp.paragraph_format.left_indent = Inches(0.24 * (indent_level + 1))
            lp.paragraph_format.space_before = Pt(0)
            lp.paragraph_format.space_after = Pt(2.5)
            lp.paragraph_format.line_spacing = 1.3
            
            num_r = lp.add_run(f"{num_prefix} ")
            format_run(num_r, font_size=Pt(10.5), color=RGBColor(37, 99, 235), bold=True)
            
            add_inline_markdown_runs(lp, item_text, default_font_size=Pt(10.5), default_color=RGBColor(30, 41, 59))
            continue

        # ==========================================================
        # 8. 普通正文段落 (五号 10.5pt，行距 1.35，优雅深石板色)
        # ==========================================================
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4.5)
        p.paragraph_format.line_spacing = 1.35
        add_inline_markdown_runs(p, stripped, default_font_size=Pt(10.5), default_color=RGBColor(30, 41, 59))

    # 收尾未闭合的 buffer
    flush_quote()
    flush_table()

    # 文章末尾页脚水印与免责提示
    end_p = doc.add_paragraph()
    end_p.paragraph_format.space_before = Pt(16)
    end_p.paragraph_format.space_after = Pt(8)
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_r = end_p.add_run(BRAND_FOOTER_NOTE)
    format_run(end_r, font_size=Pt(8.5), color=RGBColor(148, 163, 184), italic=True)

def setup_document_styles(doc: Document):
    """全局初始化 Word 文档样式体系 (中文字体锁定微软雅黑，西文字体锁定 Arial，配置全局页脚)"""
    try:
        # 设置页面边距 (标准适中：上1寸，下1寸，左右1.1寸)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.1)
            section.right_margin = Inches(1.1)

            # 配置页面底端居中水印页脚
            footer = section.footer
            footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_p.text = f"📖 微信公众号【{BRAND_OFFICIAL_ACCOUNT}】整理排版 · 仅供个人离线学习交流"
            footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if footer_p.runs:
                format_run(footer_p.runs[0], font_size=Pt(8.5), color=RGBColor(148, 163, 184), italic=True)

        # 全局 Normal 样式
        style_normal = doc.styles['Normal']
        font = style_normal.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(10.5)
        font.color.rgb = RGBColor(30, 41, 59)
        style_normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'Microsoft YaHei')
    except Exception:
        pass

class DocxExporter(BaseExporter):
    """Word (.docx) 单文件合并导出器 (出版级排版、表格原生化、中文字体严密对齐)"""

    async def export(self, articles: List[ArticleItem], filename_prefix: str) -> Path:
        output_file = self.output_dir / f"{filename_prefix}.docx"
        now_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        def _render_all_docx_sync() -> Path:
            doc = Document()
            setup_document_styles(doc)

            # 1. 封面与总标题
            if self.platform in ["微信公众号", "wechat"] or "公众号" in self.platform:
                doc_title = f"【{self.author_name}公众号合集】" if "公众号" not in self.author_name else f"【{self.author_name}文章合集】"
            else:
                doc_title = f"【{self.author_name}】文章知识合集"

            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_before = Pt(24)
            title_p.paragraph_format.space_after = Pt(8)
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_p.add_run(doc_title)
            format_run(title_run, font_size=Pt(22), color=RGBColor(15, 23, 42), bold=True)

            # 元数据说明
            meta_p = doc.add_paragraph()
            meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta_p.paragraph_format.space_after = Pt(12)
            meta_run = meta_p.add_run(f"排版整理：微信公众号【{BRAND_OFFICIAL_ACCOUNT}】   |   来源平台：{self.platform}   |   导出时间：{now_str}   |   文章总数：{len(articles)} 篇")
            format_run(meta_run, font_size=Pt(9.5), color=RGBColor(100, 116, 139), italic=True)

            # 渲染封面免责声明 Callout 引用框
            render_callout_box(doc, [BRAND_DISCLAIMER])

            # 2. 目录导航 (仅在多篇时生成)
            if len(articles) > 1:
                toc_heading = doc.add_paragraph()
                toc_heading.paragraph_format.space_before = Pt(12)
                toc_heading.paragraph_format.space_after = Pt(6)
                toc_r = toc_heading.add_run("📚 目录导航")
                format_run(toc_r, font_size=Pt(14), color=RGBColor(30, 41, 59), bold=True)

                for idx, art in enumerate(articles, 1):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.line_spacing = 1.2
                    num_r = p.add_run(f"{idx:02d}.  ")
                    format_run(num_r, font_size=Pt(10), color=RGBColor(2, 132, 199), bold=True)
                    
                    title_r = p.add_run(f"{art.title} ")
                    format_run(title_r, font_size=Pt(10), color=RGBColor(30, 41, 59))
                    
                    if art.publish_time:
                        time_r = p.add_run(f"({art.publish_time})")
                        format_run(time_r, font_size=Pt(8.5), color=RGBColor(148, 163, 184), italic=True)

                doc.add_page_break()

            # 3. 逐篇文章渲染
            for idx, art in enumerate(articles, 1):
                append_article_content_to_docx(doc, art, self.author_name, self.platform, embed_images=True)
                if idx < len(articles):
                    doc.add_page_break()

            doc.save(str(output_file))
            return output_file

        import asyncio
        return await asyncio.to_thread(_render_all_docx_sync)
